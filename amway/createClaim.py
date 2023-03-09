import requests
import re
import os
import docx
from datetime import datetime
import subprocess
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm, Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils.mailSend import sendEmail
from utils.docResult import writeData



ulrList = [
    'https://www.wildberries.ru/catalog/90053004/detail.aspx',
    'https://www.wildberries.ru/catalog/98529036/detail.aspx',
    'https://www.wildberries.ru/catalog/140970297/detail.aspx'
]

count = 2
fileForResult = []


def createClaim(item):

    myDoc = docx.Document()
    current_datetime = datetime.now()
    linkIndexPattern = re.findall('\d.*/', item)
    linkIndex = ''.join(linkIndexPattern).replace('/','')
    apiCard = f'https://card.wb.ru/cards/detail?spp=0&regions=80,64,83,4,38,33,70,68,69,86,75,30,40,48,1,22,66,31,71&pricemarginCoeff=1.0&reg=0&appType=1&emp=0&locale=ru&lang=ru&curr=rub&couponsGeo=12,3,18,15,21&dest=-1029256,-102269,-2162196,-1257786&nm={linkIndex}'
    resp = requests.get(apiCard)
    result = resp.json()
    brandName = result['data']['products'][0]['brand']

    """heading"""

    headind = myDoc.add_paragraph()
    headind.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    headind.add_run().add_picture('images/logo.png', width=Inches(5.75), height=Inches(1.02))

    simplestyle = myDoc.styles.add_style('SimpleText', WD_STYLE_TYPE.PARAGRAPH)
    simplestyle.font.name = 'Times New Roman'
    myDoc.add_paragraph(f'От:   	  Digital Rights Center, LLC (ООО "Центр цифровых прав")', style=simplestyle)
    myDoc.add_paragraph(f'Кому:     Администрация платформы Wildberries \n'
                        f'               142181, Московская область, г. Подольск, деревня Коледино,\n'
                        f'               Территория Индустриальный парк Коледино, д.6, стр.1', style=simplestyle)

    """"Date"""
    underlineStyle = myDoc.styles.add_style('dataText', WD_STYLE_TYPE.PARAGRAPH)
    underlineStyle.font.name = 'Times New Roman'
    date = myDoc.add_paragraph(f'Дата       ', style=underlineStyle)
    date.add_run(f'{current_datetime.day}.{current_datetime.month}.{current_datetime.year}').underline = True
    myDoc.add_paragraph(f'Кас.: 	 Нарушение исключительных прав на товарный знак $title_eng$', style=simplestyle)

    """title"""
    style = myDoc.styles.add_style('UserHeader', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    headind = myDoc.add_paragraph('', style=style)
    headind.add_run('ПРЕТЕНЗИЯ').bold = True

    """context"""
    contexStyle = myDoc.styles.add_style('ContextText', WD_STYLE_TYPE.PARAGRAPH)
    contexStyle.paragraph_format.first_line_indent = Mm(12)
    contexStyle.font.name = 'Times New Roman'
    myDoc.add_paragraph(f'Уважаемые господа,', style=simplestyle)


    myDoc.add_paragraph(f'Общество с ограниченной ответственностью "Центр цифровых прав", '
                        f'является уполномоченным представителем компании'
                        f' $owner_rus$ (далее – «Компания», «Правообладатель»).', style=contexStyle)

    myDoc.add_paragraph(f'Компания $owner_rus$ является обладателем исключительных прав на товарный знак'
                        f' $title_eng$ (далее – «Товарный знак»)  по свидетельству Российской Федерации '
                        f'№ _______ и высоко ценит свою интеллектуальную собственность и деловую репутацию,'
                        f' являющиеся важнейшими ее активами.', style=contexStyle)
    """context-docs"""
    contextDocs = myDoc.add_paragraph(f'Правоустанавливающие документы: ', style=simplestyle)
    contextDocs.add_run(f'\n$documents$')


    """context"""
    myDoc.add_paragraph(f'Настоящим сообщаем Вам, что на платформе Wildberries Правообладателем были обнаружены'
                        f' предложения к продаже продукции неустановленного происхождения,'
                        f' маркированной Товарными знаками Правообладателя:', style=contexStyle)
    myDoc.add_paragraph(f'Артикул товара: __________', style=simplestyle)

    contexOrigin = myDoc.add_paragraph(f'Основаниями для обращения послужили следующие условия: по вышеуказанной ссылке '
                        f'предлагается к продаже контрафакт (товар ложно обозначен товарным знаком, '
                        f'выдается за оригинальный, но не является им)/ на товаре или упаковке используется'
                        f' сходное обозначение/ товарный знак или сходное обозначение используется в названии '
                        f'карточки товара или тексте описания/товарный знак или сходное обозначение'
                        f' используется на изображении в карточке товара.', style=contexStyle)
    contexOrigin.add_run(f'\n           Источник происхождения вышеуказанной продукции Правообладателем не установлен.').bold = True

    myDoc.add_picture('images/logo.png', width=Inches(5.75), height=Inches(1.02))

    # myDoc.add_paragraph(f"1. \"{item}\".")
    # myDoc.add_paragraph(f'Расчет суммы возмещения ущерба.')
    myDoc.add_paragraph(f'В соответствии с положениями Гражданского кодекса РФ, в отсутствие доказательств'
                        f' исчерпания прав на товарный знак (т.е. документов, подтверждающих изначальное '
                        f'приобретение продукции у правообладателя или его официального представителя на '
                        f'территории РФ) презюмируется, что товарные знаки используются '
                        f'в отношении товара незаконно.', style=contexStyle)

    trademarks = myDoc.add_paragraph(f'', style=contexStyle)
    trademarks.add_run(f'Использование товарных знаков ').bold = True
    trademarks.add_run(f'в отношении товаров, введенных в гражданский оборот на территории РФ ')
    trademarks.add_run(f'без согласия правообладателя запрещено. ').bold = True
    trademarks.add_run(f'Нарушением считается не только продажа товара, но и сам факт предложения товара к '
                       f'продаже с использованием чужого товарного знака. Использование  товарного знака '
                       f'без согласия Правообладателя нарушает его исключительные права; товары, реализация '
                       f'которых сопряжена с нарушением исключительных прав, признаются контрафактными '
                       f'(ст. 1252 ГК РФ).\n')
    trademarks.add_run('              В соответствии с п.3 ст.1484 ГК РФ никто')
    trademarks.add_run(f'не вправе использовать').bold = True
    trademarks.add_run(f'без разрешения правообладателя').normal = True
    trademarks.add_run(f'сходные с его товарным знаком обозначения').bold = True
    trademarks.add_run(f'в отношении товаров, для индивидуализации которых товарный знак зарегистрирован, '
                     f'или однородных товаров, если в результате такого использования возникнет вероятность смешения.'
                     f' Отсутствие запрета не считается согласием или разрешением на такое использование.\n')

    trademarks.add_run(f'             В соответствии с законодательством РФ за нарушение прав на товарный знак '
                                         f'предусмотрены:')
    trademarks.add_run(f'гражданско-правовая ответственность, административная и '
                           f'уголовная ответственность.\n').underline= True
    trademarks.add_run(f'             В соответствии с законодательством РФ лицо, исключительное право которого нарушено, '
                        f'может требовать полного возмещения причиненных ему убытков, либо компенсации до'
                        f' 5 000 000 рублей.')

    defendant = myDoc.add_paragraph(f'Как разъяснено в пункте 12 постановления Пленума Верховного Суда '
                                    f'Российской Федерации от 23.06.2015 № 25 «О применении судами некоторых '
                                    f'положений раздела I части первой Гражданского кодекса Российской Федерации» '
                                    f'(далее – постановление от 23.06.2015 № 25), по делам о возмещении '
                                    f'убытков истец обязан доказать, что ', style=contexStyle)
    defendant.add_run(f'ответчик является лицом, в результате действий (бездействия) '
                      f'которого возник ущерб').bold = True
    defendant.add_run(f' а также факты нарушения обязательства или причинения вреда, наличие убытков.\n')
    defendant.add_run(f'              Согласно информации, полученной из открытых источников '
                      f'ООО «Вайлдберриз» не является лицом, продающим товар, однако исходя из анализа судебной '
                      f'практики ООО «Вайлдберриз» выступает в качестве ')
    defendant.add_run(f'информационного посредника. \n').bold = True
    defendant.add_run(f'              В соответствии с п. 1 ст. 1253.1. ГК РФ «Особенности '
                      f'ответственности информационного посредника» «Лицо, осуществляющее передачу материала '
                      f'в информационно - телекоммуникационной сети, в том числе в сети "Интернет", лицо, '
                      f'предоставляющее возможность размещения материала или информации, необходимой для его '
                      f'получения с использованием информационно-телекоммуникационной сети, лицо,')
    defendant.add_run().add_picture('images/logo.png', width=Inches(5.75), height=Inches(1.02))

    defendant.add_run(f' предоставляющее возможность доступа к материалу в этой сети, - информационный посредник - ')
    defendant.add_run(f'несет ответственность за нарушение интеллектуальных прав в '
                                 f'информационно-телекоммуникационной сети на общих основаниях».').bold = True


    myDoc.add_paragraph(f'Основанием для освобождения от ответственности признается отсутствие умысла со '
                        f'стороны администратора сайта и своевременное реагирование на жалобу третьего лица.',
                        style=contexStyle)

    """Claims"""
    claims = myDoc.add_paragraph(f'            В связи со всем вышесказанным, ', style=simplestyle)
    claims.add_run(f'ПРОСИМ:').bold = True

    claimsRight = myDoc.styles.add_style('ClaimsRight', WD_STYLE_TYPE.PARAGRAPH)
    claimsRight.paragraph_format.first_line_indent = Mm(20)
    claimsRight.font.name = 'Times New Roman'
    claimsRights = myDoc.add_paragraph(f'1.', style=claimsRight)
    claimsRights.add_run('  оказать нам содействие в проверке сведений в отношении указанных\n'
                         '                          товаров путем запроса у соответствующего продавца документов,\n'
                         '                          подтверждающих легальность происхождения продукции (договоры\n'
                         '                          с официальными дистрибьюторами продукции Правообладателя/\n '
                         '                         товарные накладные/платежные поручения/счета-фактуры и т.д.).\n')
    claimsRights.add_run(f'                     2.  в том случае, если данные доказательства не будут представлены\n'
                         f'                          продавцом по запросу ООО «Вайлдберриз», просим обеспечить удаление\n'
                         f'                          вышеназванных предложений о товаре с платформы Wildberries.\n')
    claimsRights.add_run(f'                     3.  представить контактные данные владельцев указанных предложений о\n'
                         f'                          товаре;\n')
    claimsRights.add_run(f'                     4.  представить письменный ответ по адресу: help@drc.law в \n'
                         f'                          установленные законом сроки.')

    finStyle = myDoc.styles.add_style('FinDoc', WD_STYLE_TYPE.PARAGRAPH)
    finStyle.paragraph_format.first_line_indent = Mm(-21)
    finStyle.font.name = 'Times New Roman'
    finStyle.font.size = Pt(11)
    finStyle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    finDoc = myDoc.add_paragraph(f'Директор ООО "ЦЦП"                     ', style=finStyle)
    finDoc.add_run(f'                                                      Ялкович В.С.\n')
    finDoc.add_run().add_picture('images/seal.png', width=Inches(1.2), height=Inches(1.22))

    """ create docx"""
    myDoc.save(f"email/ClaimFile{linkIndex}.docx")
    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', f'email/ClaimFile{linkIndex}.docx'])


    subprocess.call(f'mv ClaimFile{linkIndex}.pdf  pdf/', shell=True)

    with open('output.txt', 'w') as data:
        data.write(f'ClaimFile{linkIndex}.pdf')

    fileForResult.append( f'ClaimFile{linkIndex}.pdf')

    """call pytest how subprocess"""
    # subprocess.call('python3 -m pytest -s -v test_ordered.py --headed', shell=True)
    subprocess.call('python3 -m pytest -s -v test_appeal.py --headed', shell=True)




    # sendEmail(f'ClaimFile{linkIndex}.docx')


    print(brandName)


for i in ulrList:
    createClaim(i)
    file =''.join(fileForResult)
    data = {
        'count':count,
        'file':file,
        'link':i
    }
    with open("resultData.json", "a") as reader:
        reader.write(f'{data},\n')

    count += 1
    fileForResult.clear()

"""record resultdata"""
writeData()
subprocess.call(f'rm resultData.json', shell=True)

'https://www.ozon.ru/api/entrypoint-api.bx/page/json/v2?url=%2Fproduct%2Famway-l-o-c-zhidkost-dlya-mytya-stekol-loc-amvey-sredstvo-dlya-stekla-lok-amvey-699409218%2F%3Flayout_container%3DpdpPage2column%26layout_page_index%3D2%26sh%3DioWT5K2JUg'
'https://www.ozon.ru/api/entrypoint-api.bx/page/json/v2?url=%2Fproduct%2Fpyatnovyvoditel-amvey-sprey-dlya-predvaritelnogo-vyvedeniya-pyaten-amway-sa8-prewash-spray-400-837491259%2F%3Flayout_container%3DpdpPage2column%26layout_page_index%3D2%26sh%3DioWT5OXGjg'